"""Generate the Word version of the FirstService Residential briefing.

Reads structured slide content (mirroring TALKING_POINTS.md) and emits a
.docx with proper styles — cover page, headings, beats as bullets,
section breaks, and presenter notes.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

import sys
from datetime import datetime

_BASE = Path(__file__).resolve().parents[1] / "FirstService-Residential-Briefing-2026.docx"


def _resolve_output() -> Path:
    """Return a writable path. If the canonical filename is locked (Word has
    it open), append a timestamp so the build still succeeds."""
    if not _BASE.exists():
        return _BASE
    try:
        with open(_BASE, "ab"):
            pass
        return _BASE
    except PermissionError:
        ts = datetime.now().strftime("%Y%m%d-%H%M")
        alt = _BASE.with_name(f"{_BASE.stem}-{ts}.docx")
        print(f"[notice] {_BASE.name} is locked — writing {alt.name} instead.",
              file=sys.stderr)
        return alt


OUT = _resolve_output()

AON_RED = RGBColor(0xEB, 0x00, 0x17)
AON_INK = RGBColor(0x26, 0x28, 0x36)
AON_STONE = RGBColor(0x5D, 0x6D, 0x78)
AON_NAVY = RGBColor(0x10, 0x1E, 0x7F)


# ---------- Slide content ----------

slides: list[dict] = [
    {
        "n": 1, "title": "Cover",
        "subtitle": "The Florida Condo Market in 2026.",
        "open": "Thank you for the room. Let's get straight to it.",
        "script": (
            "This isn't a pitch deck. It's a market briefing — the same one we'd run for a "
            "publicly traded REIT, scaled to your portfolio. Florida condos are the most "
            "underwriter-watched asset class in the country right now. Every carrier on your "
            "programs — wind, casualty, D&O — is repricing your buildings. Over the next forty "
            "minutes we'll show you what they see, what they reward, and what you can actually "
            "do about it. By the end you'll have a framework you can take to your boards tomorrow."
        ),
        "beats": [
            "Briefing, not a pitch.",
            "Every line on a Florida condo is being repriced this year.",
            "You'll leave with a framework, not a sales pamphlet.",
        ],
    },
    {
        "n": 2, "title": "Your Aon team",
        "subtitle": "Four lanes of risk. One team.",
        "open": "Before any slides — meet the people who'd actually pick up the phone.",
        "script": (
            "You've got four of us in the room because condos in Florida don't fit one specialty. "
            "Kali leads our Florida property practice out of Tampa — she sees every wind placement "
            "in the state. Norbert runs Middle Market and lives in the casualty and D&O conversation. "
            "Sam is the Senior Account Executive who'd run your day-to-day. I lead Real Estate for "
            "South Florida. The reason we travel as four is simple: when something goes wrong, you "
            "don't want a generalist — you want the specialist who's already seen it twice this year."
        ),
        "beats": [
            "Four specialists, one client team.",
            "Each name maps to a specific real-world question (wind, casualty, D&O, daily ops).",
            "\"When something goes wrong, you want the specialist, not the generalist.\"",
        ],
    },
    {
        "n": 3, "title": "Agenda",
        "subtitle": "Today's runway.",
        "open": "Seven acts, forty minutes, then your floor.",
        "script": (
            "Here's the runway. We start with who Aon is and why that matters for a portfolio like "
            "yours. We move into the 2026 market — property, casualty, and D&O. We zoom into your "
            "buildings on a live risk map. Then the framework — what carriers reward, penalize, and "
            "ignore — which is the part you'll quote at your next board meeting. We leave fifteen "
            "minutes for Q&A, and we close. If anything is dragging or you want to skip ahead, just "
            "say so — these slides are clickable."
        ),
        "beats": [
            "Seven sections, ~40 minutes, 15 for Q&A.",
            "Tap any line to jump — built for the room, not for a script.",
            "The framework slide is the one to remember.",
        ],
    },
    {
        "n": 4, "title": "Aon at a glance",
        "subtitle": "The world's leading property broker.",
        "open": "Three numbers. That's the firm.",
        "script": (
            "Thirteen billion in property premium placed each year. Eight trillion — with a T — in "
            "coverage limits in force. Two hundred property brokers across the U.S., London, and "
            "Bermuda working off a single book. The reason that matters in this room: when a Bermuda "
            "reinsurer changes appetite on January 1st, we know before the press release. We're not "
            "a regional broker calling around — we're the firm that capacity follows."
        ),
        "beats": [
            "$13B premium · $8T limits · 200+ brokers.",
            "We aren't reacting to capacity — capacity is reacting to us.",
            "One credit profile across U.S. / London / Bermuda.",
        ],
    },
    {
        "n": 5, "title": "Real estate scale",
        "subtitle": "Aon represents the largest commercial RE portfolios in the U.S.",
        "open": "If the numbers on the last slide felt abstract, here's the real-estate cut.",
        "script": (
            "Three hundred billion dollars of insured real-estate values flow through Aon. Thirty "
            "percent of the largest U.S. real-estate owners — Blackstone, Brookfield, the names you "
            "read in the Wall Street Journal — name Aon as their broker of record. Three "
            "hundred-plus real-estate clients, $700 million in annual real-estate GWP, $33 billion "
            "in program limits. The point isn't bragging — it's that the same forty specialists who "
            "price a Blackstone tower price a FirstService building. Same desk, same carriers, same data."
        ),
        "beats": [
            "$300B insured RE values · 30% of the largest U.S. owners.",
            "Same desk that prices REITs prices your buildings.",
            "Forty specialists, fully aligned to RE.",
        ],
    },
    {
        "n": 6, "title": "National property broking",
        "subtitle": "We don't place coverage — we aggregate capacity.",
        "open": "Watch the globe — that's not animation, that's our placement footprint.",
        "script": (
            "Most brokers place coverage. They take your submission, send it to two or three "
            "carriers, and come back with a quote. Aon aggregates capacity — meaning when your "
            "building goes to market, it's hitting U.S. admitted, U.S. excess and surplus, London, "
            "and Bermuda simultaneously, with two hundred brokers competing those markets against "
            "each other on your behalf. That's why a bad layer in Florida can be backfilled "
            "overnight from London. One team, one credit profile, every market on earth."
        ),
        "beats": [
            "Placing ≠ aggregating. We do the latter.",
            "Four markets, simultaneous — not sequential.",
            "A bad layer in Miami can be solved in London on the same call.",
        ],
    },
    {
        "n": 7, "title": "Markets we touch",
        "subtitle": "Every market that prices your buildings.",
        "open": "These are the names on your policy — every one of them.",
        "script": (
            "Fifty-six named markets across four regions. U.S. admitted — AIG, Chubb, FM Global, "
            "Zurich. U.S. excess and surplus — RSUI, Westchester, the wholesalers that take Florida "
            "wind nobody else will. London — Lloyd's syndicates and the European primary carriers. "
            "Bermuda — the reinsurance backstop that supplies the last layer of every catastrophe "
            "program. When you ask \"did we go to market?\" — this is the market we mean. Not a "
            "panel of three."
        ),
        "beats": [
            "56 carriers, 4 regions.",
            "E&S is where Florida wind actually clears.",
            "\"Did we go to market?\" — yes, all of it.",
        ],
    },
    {
        "n": 8, "title": "Aon Client Treaty (ACT)",
        "subtitle": "28.5% of every London line, pre-secured.",
        "open": "Here's a facility no other broker on earth can offer you.",
        "script": (
            "ACT is an Aon-only treaty with eleven Lloyd's carriers — QBE, Liberty, Beazley — who "
            "have pre-committed 28.5% of capacity on every eligible London line before the rest of "
            "the market quotes. So when your renewal hits London, more than a quarter of the "
            "program is already locked, with stable pricing, on a single Lloyd's claims platform. "
            "$2.1 billion of GWP has run through it. 1,800 clients, 200 countries. It's why Aon's "
            "London quotes come back faster, tighter, and with fewer late surprises. And it's "
            "exclusive — not a \"preferred relationship,\" literally exclusive."
        ),
        "beats": [
            "28.5% of London capacity, pre-secured before quoting starts.",
            "Exclusive to Aon — no other broker has a comparable facility.",
            "Faster quotes, fewer late surprises, single-carrier claims service.",
        ],
    },
    {
        "n": 9, "title": "Service offering",
        "subtitle": "The work we do between renewals.",
        "open": "The renewal is one day. The other 364 are this.",
        "script": (
            "Most brokers earn their commission for one week a year — the placement. We earn it for "
            "the other 51. Pre-renewal strategy starts ninety days out, not three weeks. Property "
            "engineering — submission-ready COPE data, valuations, modeling — is the difference "
            "between four carriers competing and one declining. Cat modeling, premium allocation, "
            "claims advocacy, mid-term stewardship, and the Aon Property Risk Analyzer. Click any "
            "tile and we'll go deeper — but the headline is: the renewal isn't a moment, it's a workflow."
        ),
        "beats": [
            "One-week brokers vs. 51-week brokers.",
            "The data you submit decides whether four carriers compete or one declines.",
            "Click any tile in the room — these aren't slogans, they're SLAs.",
        ],
    },
    {
        "n": 10, "title": "Your day-to-day",
        "subtitle": "Background service, not a daily headache.",
        "open": "This is the slide you'll actually feel on a Tuesday.",
        "script": (
            "The renewal numbers are the ones boards remember. The Tuesday certificate request is "
            "the one that consumes your week. The Certificate Center handles COIs same-day, vendor "
            "compliance tracking is automated, and routine endorsements run on a two-business-day "
            "SLA. You get one named contact who knows your property — not a 1-800 queue. Behind "
            "that contact: a defined service bench — Account Executive, Account Manager, Claims "
            "Advocate, Risk Engineer — and AGRC on demand. The goal is invisible insurance. You "
            "shouldn't be thinking about us until you need us."
        ),
        "beats": [
            "Same-day COIs, automated vendor tracking, no dead inbox.",
            "One named contact, full bench behind them.",
            "\"Invisible insurance\" is the standard.",
        ],
    },
    {
        "n": 11, "title": "Aon Global Risk Consulting (AGRC)",
        "subtitle": "The risk consulting practice behind your placement.",
        "open": "Most brokers stop at the policy. We don't.",
        "script": (
            "AGRC is 1,300 risk-management consultants across 50 countries — a consulting practice "
            "embedded inside the broker. Five pillars: risk control, actuarial, risk-management "
            "outsourcing, captive management, and claims consulting. So when you need a "
            "replacement-cost valuation for a reserve study, a vendor-contract review, a captive "
            "feasibility, or claims preparation after a loss — it's already inside the relationship. "
            "You don't hire a separate consultant. Pre-loss, placement, post-loss — same platform, "
            "same people."
        ),
        "beats": [
            "1,300 consultants, 50 countries — inside the broker.",
            "Replaces what most boards hire a separate consultant for.",
            "Pre-loss, placement, post-loss — one bench.",
        ],
    },
    {
        "n": 12, "title": "Digital Property Profile",
        "subtitle": "Your buildings, in data the carrier trusts.",
        "open": "An underwriter doesn't price what you tell them — they price what you can prove.",
        "script": (
            "This is the toolkit that turns \"trust me, the roof is fine\" into evidence. Drone "
            "surveys of every roof and façade, 360 cameras through the mechanical rooms, IoT — "
            "Safehub for seismic, Hailios for hail — satellite imagery before and after a storm, "
            "mobile inspection forms that flow straight into our proprietary Risk Database. Same "
            "data set used pre-loss to underwrite, and post-loss to settle. The carrier sees what "
            "we see. That's the leverage."
        ),
        "beats": [
            "Drones, 360s, IoT, satellites — all feeding one database.",
            "Same data underwrites and settles.",
            "Evidence beats narrative — every renewal.",
        ],
    },
    {
        "n": 13, "title": "Aon Manuscript Form",
        "subtitle": "Same premium. Broader policy.",
        "open": "This is the quietest, most valuable thing we do.",
        "script": (
            "Most brokers hand you the carrier's policy form. We hand you ours. Four examples on "
            "this slide. Sublimits sized to your exposure, not a generic grid. Named-windstorm "
            "coverage with one trigger — wind, surge, hail, rain — so there's no finger-pointing "
            "the morning after. Replacement cost that flexes if the board redirects post-loss "
            "dollars. And — critically — you name the loss adjuster, not the insurer. Same premium "
            "dollar, materially better policy. When carriers complain about Aon at renewal, this is "
            "usually the slide they're complaining about."
        ),
        "beats": [
            "Aon's form, not the carrier's.",
            "One trigger for named storm = no coverage finger-pointing.",
            "You pick the adjuster — not the carrier.",
        ],
    },
    {
        "n": 14, "title": "Aon Rapid Response",
        "subtitle": "On site within 48–72 hours.",
        "open": "The morning after a hurricane is not when you want to be hiring an adjuster.",
        "script": (
            "Pre-storm, we map every building's emergency response and resilience baseline. Five "
            "days out, automated catastrophe alerts go to your team — wind, surge, hail, fifteen "
            "peril types. Storm hits, the Rapid Response coordinator activates. Inside 48 to 72 "
            "hours, we have loss mitigation consultants, remediation experts, risk engineers, "
            "forensic accountants, and construction estimators on-site — together. First thirty "
            "days, we stabilize income, document the loss, win the reserve and interim payments. "
            "Ongoing, complex-claims advocacy escalates anything that gets stuck. You're not hiring "
            "this team after the loss — they're already in the relationship."
        ),
        "beats": [
            "Pre-storm alerts → 48–72 hr boots on the ground → 30-day stabilization.",
            "Adjusters, engineers, forensic accountants — already named.",
            "Don't hire help after the storm. Have it standing by.",
        ],
    },
    {
        "n": 15, "title": "Renewal timeline",
        "subtitle": "120 days, 11 steps.",
        "open": "What looks like one day on your calendar is four months on ours.",
        "script": (
            "Here's the whole arc. January — kickoff, exposures finalized, modeling refreshed. "
            "February — submission released to market. March — underwriter roadshow in London and "
            "Bermuda. End of March, quotes come in. April, proposal to your board. Mid-April, "
            "order to bind. April 30 — inception. May, post-bind meeting and the stewardship "
            "cadence kicks off again. The reason we start in January for an April 30 renewal is "
            "the same reason a wedding planner doesn't book the venue the week before: the people "
            "you want — the best underwriters at the best carriers — are committed two months early."
        ),
        "beats": [
            "January start for April 30 inception.",
            "The good underwriters are committed two months out.",
            "Late submissions buy uncertainty pricing — every time.",
        ],
    },
    {
        "n": 16, "title": "Casualty & D&O",
        "subtitle": "Property is the headline. Casualty and D&O are the long tail.",
        "open": "The wind story dominates Florida. The rest is where boards actually get sued.",
        "script": (
            "Three columns. General liability is roughly flat — minus 0.4% on average — but 54% of "
            "accounts still saw increases. The average hides the spread. Umbrella and excess is "
            "still climbing — plus 3.4%, plus 7.3% on accounts above $250K premium — because "
            "nuclear verdicts keep arriving. Single-plaintiff outcomes now average $5 million-plus. "
            "And D&O — for boards specifically — special-assessment fights and construction-defect "
            "spillover are the top claim drivers. Aon places D&O for 26% of the Fortune 100, which "
            "means the playbook is mature. We just scale it down."
        ),
        "beats": [
            "GL flat on average, but the spread is wide.",
            "Umbrella still climbing — nuclear verdicts driving it.",
            "Special assessments + construction defect = top D&O drivers for condo boards.",
        ],
    },
    {
        "n": 17, "title": "P&C Market 2026",
        "subtitle": "The 8th straight quarter of rate decline.",
        "open": "If the last few years felt brutal, the data has turned.",
        "script": (
            "Four tabs — bear with me. Snapshot: average property rate is down 15.1% in Q1, eighth "
            "straight quarter of negative rate change. By peril: every CAT line is down — Tier I "
            "wind down 18%, all CAT down nearly 19%. By program: shared and layered programs are "
            "seeing minus 20%, single-carrier paper minus 9% — that gap is where structural "
            "decisions earn their keep. Outlook: continuation of the trend through Q2, supported "
            "by record reinsurance capital and a benign 2025 hurricane season. The bullet to "
            "remember: it's a buyer's market for clean property risk — but it's loss-sensitive. "
            "One meaningful 2026 cat event ends the run."
        ),
        "beats": [
            "8th straight quarter of negative rate change.",
            "Layered programs are giving up the most ground — restructure to capture it.",
            "Buyer's market, but one storm flips it.",
        ],
    },
    {
        "n": 18, "title": "70 years of storms",
        "subtitle": "One basin.",
        "open": "Just watch for a moment.",
        "script": (
            "Every line on this map is a tropical storm or hurricane that crossed the Atlantic "
            "basin since 1950. Twelve of them were Category 5 strength. Your buildings sit inside "
            "the surge zones the County evacuates first — Zone A, often Zone B. This is not the "
            "brochure map. This is the map your underwriter pulls up before they quote. When a "
            "London syndicate looks at your submission, they don't see Brickell — they see this. "
            "Which is why, on the next slide, we don't argue with the math. We work with it."
        ),
        "beats": [
            "70 years of storms, one basin, 12 at Cat-5.",
            "Your buildings sit in the first-evacuation surge zones.",
            "This is the underwriter's map — not yours.",
        ],
    },
    {
        "n": 19, "title": "2026 Atlantic outlook",
        "subtitle": "Slightly below average.",
        "open": "The good news, with a caveat.",
        "script": (
            "Twelve named storms forecast — versus a 14-storm climatology. Six hurricanes versus "
            "seven. Two major hurricanes versus three. Roughly 25 to 33% below average. Three "
            "drivers: El Niño returning, possibly to super-El Niño levels by late summer; elevated "
            "wind shear that tears storms apart before they organize; cooler Atlantic sea-surface "
            "temperatures than 2024 — less fuel. So the meteorology is friendly. But — and this is "
            "the line — it only takes one landfall to reset every assumption on this slide. Andrew "
            "was a quiet season. So was Michael. Bind before June 1."
        ),
        "beats": [
            "12 / 6 / 2 — below climatology across the board.",
            "Three drivers: El Niño, wind shear, cooler SSTs.",
            "Andrew was a quiet season. Bind before June 1.",
        ],
    },
    {
        "n": 20, "title": "Predictive AI Layer",
        "subtitle": "When does the next one reach your portfolio?",
        "open": "This is live — these are the actual probabilities your portfolio is being priced against.",
        "script": (
            "The map shows a 12-month forward outlook for the Florida watch zone, drawn from five "
            "sources: the historical HURDAT2 base rate, CSU Klotzbach, NOAA CPC, the Weather "
            "Company / Atmospheric G2, and our own predictive synthesis. Click any model — the map "
            "and stats update. The number to anchor on: probability of a Cat-1 or stronger storm "
            "passing within 50 miles of Brickell in the next 12 months. Click through them and "
            "you'll see the spread — illustrative versus cited — but the synthesis converges around "
            "a real, non-trivial number. Thirty-three days to June 1. Every named storm the basin "
            "produces tightens the rate floor and pulls capacity off the table. That's not theatre "
            "— it's how underwriters reprice in real time."
        ),
        "beats": [
            "Five models, one map, real probabilities.",
            "The Brickell 50-mile / 12-month number is the one to remember.",
            "Every named storm pre-bind = more expensive renewal.",
        ],
    },
    {
        "n": 21, "title": "Your building",
        "subtitle": "Find yours on the Miami skyline.",
        "open": "Click yours. Or — if you're brave — click your neighbor's.",
        "script": (
            "Seven of your buildings, on a live 3D map of Miami. Each pin is colored by overall "
            "risk tier. Toggle \"Manager view\" and you see what you see every day — stories, "
            "units, distance to water. Toggle \"Underwriter view\" — same buildings, scored the "
            "way a property carrier prices them. FEMA zone, distance to water, year built, loss "
            "frequency. The buildings haven't changed. Only the lens has. And the lens is what "
            "determines rate. Click any pin and we'll go a layer deeper into that asset."
        ),
        "beats": [
            "Same seven buildings, two lenses — manager vs. underwriter.",
            "The lens determines rate, not the building.",
            "Click a pin — every asset has a profile.",
        ],
    },
    {
        "n": 22, "title": "Framework — Reward",
        "subtitle": "What underwriters reward.",
        "open": "Framework time. This is the slide your boards need.",
        "script": (
            "Five things carriers actively reward. Updated COPE data — current construction, "
            "occupancy, protection, exposure on every location, submitted in advance. Recent "
            "third-party valuations within 24 months. Documented loss control — water mitigation, "
            "fire suppression upgrades, roof maintenance, with photos and dates. Stable governance "
            "— long-tenured property managers, boards that document decisions. And risk-transfer "
            "maturity — boards willing to take meaningful retentions because they understand the "
            "trade-off. None of this requires capital. All of it moves rate."
        ),
        "beats": [
            "Five rewards, all under your control.",
            "COPE + valuation + loss control + governance + retention literacy.",
            "Doesn't cost capital. Moves rate.",
        ],
    },
    {
        "n": 23, "title": "Framework — Penalize",
        "subtitle": "What underwriters penalize.",
        "open": "And the same list, inverted.",
        "script": (
            "Stale or partial COPE — missing roof age, no plumbing data, no fire-system info. When "
            "a carrier doesn't have data, they assume the worst. Replacement cost gaps — values "
            "that haven't moved with construction inflation. Underinsurance is the number-one "
            "reason carriers decline a Florida condo today. High-frequency small claims, even at "
            "low severity. Litigation-prone history — past assessment fights, board lawsuits, "
            "contractor disputes — D&O underwriters weight this heavily. And last-minute "
            "submissions. If full data lands three weeks before renewal, the carrier prices "
            "uncertainty. Every one of these is fixable, and most are free."
        ),
        "beats": [
            "Five penalties, every one fixable.",
            "Missing data = worst-case pricing.",
            "Late submissions = uncertainty premium.",
        ],
    },
    {
        "n": 24, "title": "Framework — Ignore",
        "subtitle": "What underwriters ignore.",
        "open": "And — controversially — what doesn't move the needle at all.",
        "script": (
            "Building age in isolation. 1985 versus 2005 matters less than you'd think — "
            "maintenance evidence beats year of construction. Cosmetic upgrades — new lobby, new "
            "amenities, fresh paint. Beautiful, not underwriting data. HOA awards. Long broker "
            "relationships — markets price the risk, not the producer. And marketing materials — "
            "drone footage, lifestyle photos, brochures. Underwriters want SOVs and loss runs. So "
            "if you're spending board time on these — redirect it to the reward list. That's the trade."
        ),
        "beats": [
            "Age, cosmetics, awards, broker tenure, brochures — ignored.",
            "Spend that energy on the reward list instead.",
            "Be honest with boards about what doesn't matter.",
        ],
    },
    {
        "n": 25, "title": "What you control",
        "subtitle": "Stop fighting the market. Run your half of it.",
        "open": "If you remember one slide, make it this one.",
        "script": (
            "Two columns. The right column — hurricane severity, January 1st reinsurance pricing, "
            "carrier appetite shifts, statute changes, jury verdicts, macro inflation — none of it "
            "is yours to fix. Boards burn enormous energy here. The left column is yours. "
            "Submission quality and timing — 90 to 120 days out, complete data. Documented loss "
            "control. Board governance — minutes, decisions, reserve studies. Property valuations "
            "every 24 months. Vendor and contractor due diligence. Proactive data sharing with "
            "your broker. Every renewal, the boards that obsess over the right column lose ground. "
            "The boards that execute the left column get better outcomes. Every time."
        ),
        "beats": [
            "Two columns. Only one is yours.",
            "The boards that win are the boards that focus left.",
            "Tape this slide to the conference-room wall.",
        ],
    },
    {
        "n": 26, "title": "Q&A",
        "subtitle": "The floor is yours.",
        "open": "We saved the best part for last.",
        "script": (
            "This is the part we travel for. Off-record, on-record, technical, strategic — we'd "
            "rather take the hard ones. If you're skeptical of a number on this deck, ask us about "
            "it. If something is nagging at the back of the room — that's the question we want. "
            "Four of us are here for a reason: between us, there isn't a Florida condo question "
            "that doesn't have an answer. Who's first?"
        ),
        "beats": [
            "Off-record questions welcome.",
            "The hard questions are the ones we want.",
            "Don't let people leave with unspoken doubts.",
        ],
    },
    {
        "n": 27, "title": "Close",
        "subtitle": "The market will keep evolving. Stay ahead — don't react.",
        "open": "Three things, no commitment, then we let you go.",
        "script": (
            "We're not asking you to switch brokers today. We're offering three things, no "
            "strings. One — quarterly market updates: a 30-minute briefing the week the next P&C "
            "report drops, so you and your boards see what's coming. Two — board education "
            "sessions: we come to your board meeting and translate the renewal for them, anytime. "
            "Three — a renewal readiness review, 90 days before your inception. Submission "
            "strategy, not pitch. Pick any of the three, none of the three, or all three. The "
            "number is on the next page. Thank you for the time."
        ),
        "beats": [
            "No ask today. Three optional next steps.",
            "Quarterly updates · board education · renewal readiness review.",
            "\"Pick any, none, or all. We're here either way.\"",
        ],
    },
]


# ---------- Style helpers ----------

def set_run(run, *, size=None, bold=None, italic=None, color=None, font=None):
    if font is not None:
        run.font.name = font
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), font)
        rFonts.set(qn("w:hAnsi"), font)
        rFonts.set(qn("w:cs"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_para(doc, text="", *, style=None, size=None, bold=None, italic=None,
             color=None, align=None, space_before=None, space_after=None,
             font="Calibri"):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_run(run, size=size, bold=bold, italic=italic, color=color, font=font)
    return p


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "EB0017")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run(run, size=11, color=AON_INK, font="Calibri")
    return p


# ---------- Build doc ----------

def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Default normal style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = AON_INK

    # ---- Cover ----
    add_para(doc, "", space_before=120)
    add_para(doc, "AON  |  FIRSTSERVICE RESIDENTIAL",
             size=10, bold=True, color=AON_RED, font="Calibri",
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    add_para(doc, "The Florida Condo Market in 2026",
             size=36, bold=True, color=AON_INK, font="Calibri",
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_para(doc, "Talking Points & Presenter Script",
             size=18, italic=True, color=AON_STONE, font="Calibri",
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=80)
    add_horizontal_rule(doc)
    add_para(doc, "Presenter", size=10, bold=True, color=AON_STONE,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "Karim Nasser, SVP — Aon Miami",
             size=14, color=AON_INK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    add_para(doc, "Audience", size=10, bold=True, color=AON_STONE,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "FirstService Residential Property Managers",
             size=14, color=AON_INK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    add_para(doc, "Date · Location", size=10, bold=True, color=AON_STONE,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "May 6, 2026  ·  Miami, FL",
             size=14, color=AON_INK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=120)

    # ---- How to use ----
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

    add_para(doc, "How to use this document",
             size=20, bold=True, color=AON_NAVY, space_after=12)
    add_horizontal_rule(doc)
    add_para(doc,
             "Each slide gets three things:",
             size=11, color=AON_INK, space_after=6)
    add_bullet(doc, "Open — the line you walk in with.")
    add_bullet(doc, "Script — what you actually say (≈60–90 sec).")
    add_bullet(doc, "Beats — three punchy bullets in case the script slips.")
    add_para(doc, "", space_after=12)
    add_para(doc,
             "Pace target: ~75 seconds per slide. The framework slides (22–25) and the "
             "portfolio map (21) deserve more time. The carrier wall (7) and renewal "
             "timeline (15) can move quickly.",
             size=11, italic=True, color=AON_STONE, space_after=18)

    # ---- Slide entries ----
    for s in slides:
        # New page for every slide for printability
        p = doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)

        # Header line: "01 · COVER"
        header_p = doc.add_paragraph()
        header_p.paragraph_format.space_after = Pt(2)
        run = header_p.add_run(f"{s['n']:02d}  ·  ")
        set_run(run, size=11, bold=True, color=AON_RED, font="Calibri")
        run = header_p.add_run(s["title"].upper())
        set_run(run, size=11, bold=True, color=AON_STONE, font="Calibri")

        # Subtitle / slide headline
        add_para(doc, s["subtitle"],
                 size=22, bold=True, color=AON_INK, space_after=10)
        add_horizontal_rule(doc)

        # Open
        add_para(doc, "OPEN", size=9, bold=True, color=AON_RED, space_after=2)
        add_para(doc, f"“{s['open']}”",
                 size=12, italic=True, color=AON_INK, space_after=14)

        # Script
        add_para(doc, "SCRIPT", size=9, bold=True, color=AON_RED, space_after=2)
        add_para(doc, s["script"],
                 size=11, color=AON_INK, space_after=14)

        # Beats
        add_para(doc, "BEATS", size=9, bold=True, color=AON_RED, space_after=2)
        for b in s["beats"]:
            add_bullet(doc, b)

    # ---- Presenter notes appendix ----
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

    add_para(doc, "Presenter notes",
             size=22, bold=True, color=AON_NAVY, space_after=10)
    add_horizontal_rule(doc)

    notes = [
        ("Pace",
         "27 slides, ~40 minutes content, 15 minutes Q&A. ~75 seconds per slide is the "
         "rhythm. The framework slides (22–25) and the portfolio map (21) deserve more "
         "time. The carrier wall (7) and renewal timeline (15) can move quickly."),
        ("Hand-offs",
         "Kali on Florida wind capacity (slides 17–19). Norbert on casualty / D&O (slide 16). "
         "Sam on day-to-day execution (slide 10). Karim takes the rest — and the close."),
        ("The two slides that sell the room",
         "Manuscript Form (13) and What You Control (25). Slow down on both."),
        ("If running long",
         "Collapse 6 and 7 into one beat. Skip 11 (AGRC) — it lives in the appendix. "
         "Never skip the framework."),
        ("If asked “what does this cost”",
         "“Today, nothing. The renewal readiness review is on us. If we earn the "
         "placement after that, the carrier pays the commission — same as the broker you "
         "have now. The difference is what's wrapped around the placement.”"),
    ]
    for label, body in notes:
        add_para(doc, label.upper(),
                 size=10, bold=True, color=AON_RED, space_before=10, space_after=2)
        add_para(doc, body, size=11, color=AON_INK, space_after=6)

    # ---- End mark ----
    add_para(doc, "", space_before=24)
    add_para(doc, "— end —",
             size=10, italic=True, color=AON_STONE,
             align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(OUT)
    print(f"Wrote: {OUT}")


if __name__ == "__main__":
    build()
